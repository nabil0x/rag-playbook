"""Word (.docx) document loader.

Loader block: parse Word documents *structure-aware*.

A .docx is a zip of XML whose body is a sequence of blocks — paragraphs
(``<w:p>``) and tables (``<w:tbl>``) — in document order. Two ways to read it:

1. ``mode="naive"`` — flatten ``doc.paragraphs`` like a text file. Fast and
   tempting, but python-docx's ``.paragraphs`` property only returns
   *body-level* paragraphs: text inside table cells lives in a different part
   of the XML tree, so this mode silently drops every table — headers, data
   rows, everything. Good for demonstrating *why* a structure-aware loader is
   needed; bad for retrieval accuracy.

2. ``mode="structured"`` (default) — walk the body *in order* with
   ``doc.element.body.iterchildren()``, the only reliable way to see
   paragraphs and tables interleaved as they appear. Each block becomes one
   :class:`langchain_core.documents.Document`:

   * a ``Heading 1``/``Heading 2``/``Heading 3`` style paragraph → ``# Title``,
     ``## Title``, ... (so a header-aware splitter can reuse the structure)
   * any other paragraph → its text
   * a table → a **markdown table** (header row + ``| --- |`` separator +
     data rows) so the rows survive as one unit

   Every ``Document`` carries ``source``, ``type`` (``heading``/``paragraph``/
   ``table``) and ``heading`` — the nearest preceding heading, omitted when
   there is none.

Note: the file is named ``word.py`` (not ``docx.py``) on purpose — a module
called ``docx`` would shadow the ``python-docx`` package itself when the
``src/loaders/`` directory is on ``sys.path``.

See Topics/Project-13-Multi-format-RAG/README.md (multi-format loaders).
"""

from __future__ import annotations

import os
from typing import Optional

from langchain_core.documents import Document


class WordLoader:
    """Load text and tables from a Word document.

    Args:
        path: path to the .docx file.
        mode: ``"structured"`` (default) walks the body in document order and
            re-serializes tables as markdown; ``"naive"`` flattens only
            ``doc.paragraphs`` and drops all tables (for comparison).
    """

    def __init__(self, path: str, mode: str = "structured"):
        if mode not in ("structured", "naive"):
            raise ValueError(f"mode must be 'structured' or 'naive', got {mode!r}")
        self.path = path
        self.mode = mode

    # ------------------------------------------------------------------ #
    # internals
    # ------------------------------------------------------------------ #
    @staticmethod
    def _style_name(paragraph) -> Optional[str]:
        """Return the paragraph's style name, or None when unreadable."""
        try:
            return (paragraph.style.name or "").strip() or None
        except Exception:
            return None

    @staticmethod
    def _heading_level(style_name: Optional[str]) -> Optional[int]:
        """'Heading 1'/'Heading 2'/'Heading 3' style → 1/2/3, else None."""
        if not style_name or not style_name.lower().startswith("heading"):
            return None
        try:
            return int(style_name.split()[-1])
        except ValueError:
            return None

    @staticmethod
    def _table_to_markdown(table) -> str:
        """Render a python-docx table as a markdown table.

        First row = header, second line = ``| --- |`` separators, then the
        data rows. Multi-line cell text is flattened onto one line so the
        markdown stays valid.
        """
        lines = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        return "\n".join(lines)

    # ------------------------------------------------------------------ #
    # public API
    # ------------------------------------------------------------------ #
    def load(self) -> list[Document]:
        """Parse the .docx into a list of langchain Documents."""
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError(
                "WordLoader needs python-docx: pip install python-docx"
            ) from exc
        if not os.path.exists(self.path):
            raise FileNotFoundError(f"Word file not found: {self.path}")

        doc = DocxDocument(self.path)

        if self.mode == "naive":
            # Flatten body-level paragraphs only — tables are dropped.
            return [
                Document(
                    page_content=para.text.strip(),
                    metadata={"source": self.path, "type": "paragraph"},
                )
                for para in doc.paragraphs
                if para.text.strip()
            ]

        # mode == "structured": walk the body in document order.
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        documents: list[Document] = []
        heading: Optional[str] = None

        for child in doc.element.body.iterchildren():
            if child.tag.endswith("}p"):
                para = Paragraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue
                style = self._style_name(para)
                level = self._heading_level(style)
                if level:
                    documents.append(
                        Document(
                            page_content="#" * level + " " + text,
                            metadata={
                                "source": self.path,
                                "type": "heading",
                                "heading": text,
                            },
                        )
                    )
                    heading = text
                else:
                    meta = {"source": self.path, "type": "paragraph"}
                    if heading:
                        meta["heading"] = heading
                    documents.append(
                        Document(page_content=text, metadata=meta)
                    )
            elif child.tag.endswith("}tbl"):
                meta = {"source": self.path, "type": "table"}
                if heading:
                    meta["heading"] = heading
                documents.append(
                    Document(
                        page_content=self._table_to_markdown(Table(child, doc)),
                        metadata=meta,
                    )
                )

        return documents


if __name__ == "__main__":
    import sys

    try:
        path = (
            sys.argv[1]
            if len(sys.argv) > 1
            else "Data/SD-01-word/fcc-nationwide-eas-test-2021.docx"
        )
        for mode in ("naive", "structured"):
            docs = WordLoader(path, mode=mode).load()
            types: dict[str, int] = {}
            for d in docs:
                t = d.metadata.get("type", "?")
                types[t] = types.get(t, 0) + 1
            print(f"[{mode}] {len(docs)} documents, by type: {types}")
        print("\nFirst structured unit:")
        print(WordLoader(path).load()[0].page_content[:200])
    except (ImportError, FileNotFoundError) as exc:
        print(f"SKIP: {exc}")
